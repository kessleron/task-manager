from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory
import re
import uuid
import os
from docx import Document
from bs4 import BeautifulSoup
import tempfile
import xml.etree.ElementTree as ET

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'  # נדרש עבור הודעות flash

# יצירת תיקיית העלאות זמנית
UPLOAD_FOLDER = tempfile.mkdtemp()
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # מגבלת גודל קובץ - 16 מגה

# מאגר משימות "בזיכרון" (לדוגמה פשוטה, במערכת אמיתית נשתמש במסד נתונים)
tasks_db = []
# משימות פוטנציאליות שחולצו
potential_tasks_session = []

# שם הקובץ האחרון שהועלה - שימושי להצגת פרטי הקובץ בממשק
last_uploaded_file = None

# דפוסי regex פשוטים לזיהוי משימות ותאריכים בעברית
TASK_PATTERNS = [
    re.compile(r"(?:צריך ל|יש ל|לבצע|להכין|לבדוק|לסיים|לתאם|תזכורת:|משימה:)\s*(.*?)(?:\s*(?:עד|ב-)\s*([\w\s\d.:/-]+))?(?:\n|$)"),
    re.compile(r"(?:בבקשה\s*)?([תפ]על\w*)\s+(.*?)(?:\s*(?:עד|ב-)\s*([\w\s\d.:/-]+))?(?:\n|$)")  # לטיפול בפעלים בציווי
]

def extract_tasks_from_text(text):
    """חילוץ משימות מטקסט רגיל באמצעות ביטויים רגולריים"""
    extracted = []
    lines = text.splitlines()
    for line in lines:
        if not line.strip():
            continue
        for pattern in TASK_PATTERNS:
            match = pattern.search(line)
            if match:
                # מקרה 1: (מילת מפתח) (תיאור) (עד תאריך)
                if len(match.groups()) == 2 and match.group(1):
                    description = match.group(1).strip()
                    due_date = match.group(2).strip() if match.group(2) else ""
                    extracted.append({"id": str(uuid.uuid4()), "description": description, "due_date": due_date, "raw_text": line})
                    break  # מצאנו התאמה, נעבור לשורה הבאה
                # מקרה 2: (פועל בציווי) (תיאור) (עד תאריך)
                elif len(match.groups()) == 3 and match.group(2):
                    action_verb = match.group(1).strip()
                    description_body = match.group(2).strip()
                    description = f"{action_verb} {description_body}"
                    due_date = match.group(3).strip() if match.group(3) else ""
                    extracted.append({"id": str(uuid.uuid4()), "description": description, "due_date": due_date, "raw_text": line})
                    break  # מצאנו התאמה, נעבור לשורה הבאה
    return extracted

def find_summary_section_from_paragraphs(doc):
    """מחפש את קטע הסיכום בפסקאות של מסמך Word"""
    found_summary_section = False
    summary_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # בדיקה אם מצאנו את כותרת הסיכום
        if re.search(r'סיכום.*המשך טיפול|המשך טיפול.*סיכום', text, re.IGNORECASE):
            found_summary_section = True
            continue
            
        # אם כבר נמצאנו בקטע הסיכום, נשמור את התוכן
        if found_summary_section and text:
            summary_content.append(text)
            
    return summary_content

def extract_tasks_from_table(table):
    """חילוץ משימות מטבלה במסמך Word"""
    tasks = []
    
    # בדיקה שיש לפחות 2 שורות בטבלה (שורת כותרת + לפחות שורת משימה אחת)
    if len(table.rows) < 2:
        return tasks
    
    # איתור האינדקסים של העמודות הרלוונטיות
    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
    
    task_col_idx = -1
    responsible_col_idx = -1
    deadline_col_idx = -1
    
    for idx, header in enumerate(header_cells):
        if any(term in header.lower() for term in ['החלטה', 'המשך טיפול', 'משימה']):
            task_col_idx = idx
        elif any(term in header.lower() for term in ['אחראי', 'באחריות']):
            responsible_col_idx = idx
        elif any(term in header.lower() for term in ['לוח זמנים', 'תאריך', 'מועד', 'deadline']):
            deadline_col_idx = idx
    
    # אם לא מצאנו עמודת משימות, נחזיר רשימה ריקה
    if task_col_idx == -1:
        return tasks
    
    # עיבוד כל שורה מהשורה השנייה ואילך (דילוג על שורת הכותרת)
    for row_idx in range(1, len(table.rows)):
        row = table.rows[row_idx]
        
        # וידוא שיש מספיק תאים בשורה
        if len(row.cells) <= task_col_idx:
            continue
            
        task_text = row.cells[task_col_idx].text.strip()
        
        # אם אין טקסט משימה, נדלג
        if not task_text:
            continue
            
        # חילוץ אחראי ומועד אם הם קיימים
        responsible = row.cells[responsible_col_idx].text.strip() if responsible_col_idx != -1 and responsible_col_idx < len(row.cells) else ""
        deadline = row.cells[deadline_col_idx].text.strip() if deadline_col_idx != -1 and deadline_col_idx < len(row.cells) else ""
        
        # בניית תיאור המשימה
        if responsible:
            description = f"{task_text} - אחראי: {responsible}"
        else:
            description = task_text
            
        raw_text = description
        if deadline:
            raw_text += f" (לו\"ז: {deadline})"
            
        tasks.append({
            "id": str(uuid.uuid4()),
            "description": description,
            "due_date": deadline,
            "raw_text": raw_text
        })
        
    return tasks

def find_summary_tables(doc):
    """מחפש טבלאות סיכום במסמך Word"""
    tasks = []
    
    # חיפוש בפסקאות לפני טבלאות
    summary_section_found = False
    
    for para in doc.paragraphs:
        if re.search(r'סיכום.*המשך טיפול|המשך טיפול.*סיכום', para.text, re.IGNORECASE):
            summary_section_found = True
            break
    
    # אם נמצא סעיף סיכום, נחפש טבלאות שמופיעות אחריו
    if summary_section_found:
        found_table = False
        
        for table in doc.tables:
            # נחפש בשורה הראשונה של הטבלה אם יש כותרות שרלוונטיות למשימות
            if len(table.rows) > 0:
                header_row = [cell.text.strip().lower() for cell in table.rows[0].cells]
                
                has_task_column = any(any(term in cell for term in ['החלטה', 'המשך טיפול', 'משימה']) for cell in header_row)
                
                if has_task_column:
                    found_table = True
                    tasks.extend(extract_tasks_from_table(table))
                    
        if not found_table:
            # אם לא מצאנו טבלאות רלוונטיות, ננסה לחפש משימות בטקסט הרגיל
            summary_content = find_summary_section_from_paragraphs(doc)
            for text in summary_content:
                tasks.extend(extract_tasks_from_text(text))
    else:
        # אם לא מצאנו סעיף סיכום מפורש, נבדוק את כל הטבלאות במסמך
        for table in doc.tables:
            if len(table.rows) > 0:
                header_row = [cell.text.strip().lower() for cell in table.rows[0].cells]
                
                # בדיקה אם הטבלה נראית כמו טבלת משימות
                has_task_column = any(any(term in cell for term in ['החלטה', 'המשך טיפול', 'משימה']) for cell in header_row)
                has_responsible_column = any(any(term in cell for term in ['אחראי', 'באחריות']) for cell in header_row)
                
                if has_task_column and has_responsible_column:
                    tasks.extend(extract_tasks_from_table(table))
    
    return tasks

def extract_tasks_from_word_file(file_path):
    """פונקציה ראשית לחילוץ משימות מקובץ Word"""
    try:
        doc = Document(file_path)
        return find_summary_tables(doc)
    except Exception as e:
        print(f"שגיאה בקריאת קובץ Word: {str(e)}")
        return []

@app.route('/', methods=['GET', 'POST'])
def index():
    global potential_tasks_session, last_uploaded_file
    if request.method == 'POST':
        if 'text_input' in request.form:
            text = request.form['text_input']
            potential_tasks_session = extract_tasks_from_text(text)
            if not potential_tasks_session:
                flash("לא נמצאו משימות בטקסט שהוזן")
            return render_template('index.html', 
                                   potential_tasks=potential_tasks_session, 
                                   tasks=tasks_db, 
                                   last_file=last_uploaded_file)
        
        elif 'file_upload' in request.files:
            file = request.files['file_upload']
            if file and file.filename != '':
                if file.filename.endswith(('.docx')):  # רק קבצי docx נתמכים כרגע
                    try:
                        # שמירת הקובץ בתיקייה זמנית
                        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                        file.save(file_path)
                        last_uploaded_file = {
                            'name': file.filename,
                            'path': file_path
                        }
                        
                        # חילוץ המשימות מהקובץ
                        potential_tasks_session = extract_tasks_from_word_file(file_path)
                        
                        if not potential_tasks_session:
                            flash("לא נמצאו משימות בקובץ שהועלה או שהפורמט אינו תואם")
                        else:
                            flash(f"נמצאו {len(potential_tasks_session)} משימות בקובץ", 'success')
                    except Exception as e:
                        flash(f"שגיאה בעיבוד הקובץ: {str(e)}")
                        print(f"Exception: {str(e)}")
                else:
                    flash("נא להעלות קובץ Word בלבד (סיומת .docx)")
            else:
                flash("לא נבחר קובץ")
                
            return render_template('index.html', 
                                   potential_tasks=potential_tasks_session, 
                                   tasks=tasks_db, 
                                   last_file=last_uploaded_file)
        
        elif 'confirm_task_id' in request.form:  # אישור משימה ספציפית מהרשימה הפוטנציאלית
            task_id_to_confirm = request.form['confirm_task_id']
            task_to_add = next((pt for pt in potential_tasks_session if pt["id"] == task_id_to_confirm), None)
            if task_to_add:
                # נותן ID חדש למשימה ב-DB כדי למנוע כפילויות אם המשתמש ינסה לאשר שוב
                new_task = {
                    "id": str(uuid.uuid4()),
                    "description": task_to_add["description"],
                    "due_date": task_to_add["due_date"],
                    "done": False
                }
                tasks_db.append(new_task)
                # הסר את המשימה מהרשימה הפוטנציאלית לאחר אישור
                potential_tasks_session = [pt for pt in potential_tasks_session if pt["id"] != task_id_to_confirm]
                flash("המשימה נוספה בהצלחה", 'success')
            return redirect(url_for('index'))

        elif 'confirm_all_tasks' in request.form:  # אישור כל המשימות בבת אחת
            for pt_task in potential_tasks_session:
                new_task = {
                    "id": str(uuid.uuid4()),
                    "description": pt_task["description"],
                    "due_date": pt_task["due_date"],
                    "done": False
                }
                tasks_db.append(new_task)
                
            flash(f"{len(potential_tasks_session)} משימות נוספו בהצלחה", 'success')
            potential_tasks_session = []  # ניקוי רשימת המשימות הפוטנציאליות
            return redirect(url_for('index'))

        elif 'manual_description' in request.form:  # הוספת משימה ידנית
            description = request.form['manual_description']
            due_date = request.form.get('manual_due_date', '')  # .get למקרה שהשדה לא קיים
            if description:
                tasks_db.append({"id": str(uuid.uuid4()), "description": description, "due_date": due_date, "done": False})
                flash("המשימה נוספה בהצלחה", 'success')
            return redirect(url_for('index'))

    return render_template('index.html', 
                          potential_tasks=potential_tasks_session, 
                          tasks=tasks_db, 
                          last_file=last_uploaded_file)

@app.route('/toggle/<task_id>')
def toggle_task(task_id):
    for task in tasks_db:
        if task['id'] == task_id:
            task['done'] = not task['done']
            flash("סטטוס המשימה עודכן", 'success')
            break
    return redirect(url_for('index'))

@app.route('/delete/<task_id>')
def delete_task(task_id):
    global tasks_db
    tasks_db = [task for task in tasks_db if task['id'] != task_id]
    flash("המשימה נמחקה בהצלחה", 'success')
    return redirect(url_for('index'))

@app.route('/clear_potential')
def clear_potential():
    global potential_tasks_session
    potential_tasks_session = []
    flash("רשימת המשימות הפוטנציאליות נוקתה", 'success')
    return redirect(url_for('index'))

# ניקוי תיקיית ההעלאות בסיום הפעלת האפליקציה
@app.teardown_appcontext
def cleanup_upload_folder(exception):
    import shutil
    shutil.rmtree(UPLOAD_FOLDER, ignore_errors=True)

if __name__ == '__main__':
    app.run(debug=True)
